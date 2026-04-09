#!/usr/bin/env python3
"""
Sportly HTML → DOCX Converter
Ricrea il documento di progetto con struttura, colori e formattazione fedeli.

Prerequisiti:
    pip install python-docx

Uso:
    python convert_sportly.py
"""

from docx import Document
from docx.shared import (
    Pt, Cm, Mm, Emu, RGBColor, Inches
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import io

# ═══════════════════════════════════════════
#  PALETTE COLORI (corrispondenza CSS)
# ═══════════════════════════════════════════
C = {
    'c1': RGBColor(0xa3, 0xc6, 0x39),   # lime
    'c2': RGBColor(0x0d, 0x8c, 0x8c),   # teal
    'c3': RGBColor(0x0a, 0x7a, 0x7a),
    'c4': RGBColor(0x5d, 0xaa, 0x60),   # green
    'c5': RGBColor(0x8c, 0xb8, 0x5a),
    'c6': RGBColor(0x3a, 0x9e, 0x7a),
    'c7': RGBColor(0x2d, 0x9a, 0x90),
    'c8': RGBColor(0x38, 0xa8, 0x98),
    'ink':    RGBColor(0x0c, 0x0c, 0x0c),
    'ink2':   RGBColor(0x3a, 0x3a, 0x3a),
    'ink3':   RGBColor(0x71, 0x71, 0x71),
    'ink4':   RGBColor(0xa8, 0xa8, 0xa8),
    'white':  RGBColor(0xff, 0xff, 0xff),
    'red':    RGBColor(0xe0, 0x52, 0x52),
    'surface': RGBColor(0xfa, 0xfa, 0xf8),
    'surface2': RGBColor(0xf3, 0xf3, 0xef),
    'surface3': RGBColor(0xea, 0xea, 0xe4),
    'border': RGBColor(0xe2, 0xe2, 0xda),
}


# ═══════════════════════════════════════════
#  UTILITÀ
# ═══════════════════════════════════════════
def _hex(color):
    """RGBColor → hex string #RRGGBB."""
    return f'{color[0]:02x}{color[1]:02x}{color[2]:02x}'


def _cell_shading(cell, hex_color):
    """Applica sfondo a una cella."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Imposta bordi cella."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [('top', top), ('bottom', bottom),
                       ('left', left), ('right', right)]:
        if val:
            sz, color = val
            b = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" '
                f'w:sz="{sz}" w:space="0" w:color="{_hex(color)}"/>'
            )
            borders.append(b)
    tcPr.append(borders)


def _add_colored_bar(doc, label, hex_color, width_pct, value):
    """Crea una barra colorata con label e valore."""
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    # Label
    c0 = t.cell(0, 0)
    c0.text = ''
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label)
    run.font.size = Pt(7)
    run.font.color.rgb = C['ink2']
    c0.width = Cm(2.8)

    # Barra
    c1 = t.cell(0, 1)
    _cell_shading(c1, _hex(C['surface3']))
    # Indicatore interno
    p_inner = c1.paragraphs[0]
    bar_run = p_inner.add_run(' ')
    bar_run.font.size = Pt(1)
    p_inner.paragraph_format.space_before = Pt(4)
    p_inner.paragraph_format.space_after = Pt(4)
    # Overlay color bar
    bar_shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    p_inner._p.get_or_add_pPr().append(bar_shading)
    c1.width = Cm(8)

    # Valore
    c2 = t.cell(0, 2)
    c2.text = ''
    p = c2.paragraphs[0]
    run = p.add_run(value)
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run.bold = True
    c2.width = Cm(1.6)

    # Rimuovi bordi tabella
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)


def _add_bullet_list(doc, items, color=C['c1']):
    """Aggiunge una lista con punti colorati."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        # Bullet colorato
        run = p.add_run('● ')
        run.font.size = Pt(6)
        run.font.color.rgb = color
        run = p.add_run(item)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink2']


def _add_separator(doc):
    """Linea separatrice leggera."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Tabella 1×1 con bordo inferiore come linea
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    _cell_border(c, bottom=(4, C['border']))
    c.text = ' '
    p2 = c.paragraphs[0]
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run = p2.runs[0] if p2.runs else p2.add_run()
    run.font.size = Pt(1)
    run.font.color.rgb = C['white']


def _add_spacer(doc, pts=6):
    """Spazio vuoto."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(' ')
    run.font.size = Pt(1)


# ═══════════════════════════════════════════
#  BUILD PRINCIPALE
# ═══════════════════════════════════════════
def create_sportly_docx():
    doc = Document()

    # ── Impostazioni pagina A4 ──
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

    # ── Configurazione stile default ──
    style = doc.styles['Normal']
    style.font.name = 'DM Sans'
    style.font.size = Pt(8.8)
    style.font.color.rgb = C['ink2']
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5

    # ── Configurazione Heading 1 ──
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Space Grotesk'
    h1_style.font.size = Pt(26)
    h1_style.font.bold = True
    h1_style.font.color.rgb = C['ink']
    h1_style.paragraph_format.space_before = Pt(8)
    h1_style.paragraph_format.space_after = Pt(8)

    # ── Configurazione Heading 2 ──
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Instrument Serif'
    h2_style.font.size = Pt(13)
    h2_style.font.color.rgb = C['ink']
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════
    #  PAGINA 1 — COPERTINA
    # ══════════════════════════════════════

    # Strip laterale simulata con bordo sinistro
    accent_p = doc.add_paragraph()
    accent_p.paragraph_format.space_before = Pt(0)
    accent_p.paragraph_format.space_after = Pt(4)
    run = accent_p.add_run(' ')
    run.font.size = Pt(1)

    # Header logo
    _p = doc.add_paragraph()
    _p.paragraph_format.space_before = Pt(0)
    _p.paragraph_format.space_after = Pt(0)
    run = _p.add_run('█ ')
    run.font.size = Pt(14)
    run.font.color.rgb = C['ink']
    run = _p.add_run('  SPORTLY — DOCUMENTO DI PROGETTO')
    run.font.size = Pt(6)
    run.font.color.rgb = C['ink4']
    run.font.name = 'JetBrains Mono'

    _add_separator(doc)

    # Meta line
    _p = doc.add_paragraph()
    _p.paragraph_format.space_after = Pt(2)
    run = _p.add_run('PRIMA CONSEGNA')
    run.font.size = Pt(6.5)
    run.font.color.rgb = C['c2']
    run.font.name = 'JetBrains Mono'
    run.bold = True
    run.font.underline = True
    run = _p.add_run('    Classe 5I — Aprile 2025')
    run.font.size = Pt(6.5)
    run.font.color.rgb = C['ink3']
    run.font.name = 'JetBrains Mono'

    _add_spacer(doc, 8)

    # Titolo grande
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('Sport')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(52)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run('ly')
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(52)
    run.font.italic = True
    run.font.color.rgb = C['c2']
    run = p.add_run('.')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(52)
    run.font.bold = True
    run.font.color.rgb = C['ink']

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('Project')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(52)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run('Blueprint.')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(52)
    run.font.bold = True
    run.font.color.rgb = C['ink']

    # Sottotitolo
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        'Piattaforma digitale per la prenotazione di attività sportive e ricreative '
        '— dal campo da padel alla lezione di yoga, tutto in un ecosistema integrato '
        'che trasforma le associazioni sportive locali in realtà digitali competitive.'
    )
    run.font.size = Pt(9.5)
    run.font.light = True
    run.font.color.rgb = C['ink3']

    _add_separator(doc)
    _add_spacer(doc, 4)

    # KPI row
    kpis = [
        ('Mercato Non Digitale', '65%',
         'ASD italiane ancora su sistemi analogici'),
        ('Target Primario', '18–45',
         'Sportivi attivi con smartphone connesso'),
        ('Proiezione Mese 12', '15K€',
         'Ricavo mensile stimato a regime'),
    ]

    t = doc.add_table(rows=3, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for col_i, (label, val, desc) in enumerate(kpis):
        # Label
        c = t.cell(0, col_i)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(6)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['ink3']
        p.paragraph_format.space_after = Pt(2)

        # Valore
        c = t.cell(1, col_i)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Space Grotesk'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = C['c2']
        p.paragraph_format.space_after = Pt(2)

        # Descrizione
        c = t.cell(2, col_i)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(desc)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink3']

        if col_i < 2:
            _cell_border(t.cell(0, col_i), right=(4, C['border']))
            _cell_border(t.cell(1, col_i), right=(4, C['border']))
            _cell_border(t.cell(2, col_i), right=(4, C['border']))

    _add_spacer(doc, 10)

    # Donut chart description
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('Analisi di Mercato — Penetrazione Digitale ASD')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _p = doc.add_paragraph()
    run = _p.add_run('■ ')
    run.font.size = Pt(8)
    run.font.color.rgb = C['c1']
    run = _p.add_run('65% strutture non digitalizzate — opportunità primaria')
    run.font.size = Pt(7.5)
    run.font.color.rgb = C['ink2']
    _p = doc.add_paragraph()
    run = _p.add_run('■ ')
    run.font.size = Pt(8)
    run.font.color.rgb = C['c2']
    run = _p.add_run('35% già digitalizzate — mercato concorrenziale')
    run.font.size = Pt(7.5)
    run.font.color.rgb = C['ink2']

    _p = doc.add_paragraph()
    run = _p.add_run('Stima su ASD dilettantistiche italiane — CONI / ISTAT 2024')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_spacer(doc, 8)

    # Info grid
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    info = [
        ('Tipologia di Progetto', 'Startup Digitale',
         'App Mobile (iOS + Android) + Portale Web'),
        ('Modello di Business', 'Multi-Revenue',
         'SaaS B2B + Commissioni + Freemium'),
    ]
    for i, (lbl, val, sub) in enumerate(info):
        for row_offset, text, sz, clr, bold in [
            (0, lbl, Pt(6), C['ink3'], False),
            (1, val, Pt(11), C['ink'], True),
            (2, sub, Pt(7.5), C['ink3'], False),
        ]:
            c = t.cell(i * 3 + row_offset, 0)
            c.text = ''
            p = c.paragraphs[0]
            run = p.add_run(text)
            run.font.size = sz
            run.font.color.rgb = clr
            if bold:
                run.font.name = 'Instrument Serif'
            if row_offset == 0:
                run.font.name = 'JetBrains Mono'
                run.font.text_transform = True  # uppercase sim
            p.paragraph_format.space_after = Pt(1)

    # Sfondo celle info
    for row in range(6):
        for col in range(2):
            _cell_shading(t.cell(row, col), _hex(C['surface2']))

    _add_spacer(doc, 6)

    # Team strip
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    c = t.cell(0, 0)
    c.text = ''
    _cell_shading(c, _hex(C['c2']))
    p = c.paragraphs[0]
    run = p.add_run('TEAM')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['white']
    run.bold = True
    c.width = Cm(1.5)

    c = t.cell(0, 1)
    c.text = ''
    _cell_border(c, top=(4, C['c2']), bottom=(4, C['border']),
                 left=(4, C['c2']), right=(4, C['border']))
    p = c.paragraphs[0]
    run = p.add_run('Team 6-3-14')
    run.font.size = Pt(12)
    run.font.name = 'Space Grotesk'
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p = c.add_paragraph()
    run = p.add_run(
        'Gestione Progetto e Organizzazione di Impresa — '
        'Classe 5I · Consegna: 10 Aprile 2025'
    )
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']

    # ══════════════════════════════════════
    #  PAGINA 2 — SCELTA DEL PROGETTO
    # ══════════════════════════════════════
    doc.add_page_break()

    # Running header
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('█ SPORTLY')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run = p.add_run('    01 — Scelta del Progetto')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']
    run = p.add_run('                                        02 / 09')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_separator(doc)

    # Eyebrow
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SEZIONE 01')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('Scelta del')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run('Progetto')
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(26)
    run.font.italic = True
    run.font.color.rgb = C['c2']

    # Body text
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('Il team ha selezionato lo sviluppo di ')
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']
    run = p.add_run('Sportly')
    run.font.size = Pt(8.8)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run(
        ', applicazione mobile e web per la prenotazione di attività sportive '
        'e ricreative. La scelta nasce dall\'intersezione di due dinamiche reali '
        'nel panorama sportivo italiano: la '
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']
    run = p.add_run('crescente domanda di servizi digitali')
    run.font.size = Pt(8.8)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run(' da parte degli sportivi e la ')
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']
    run = p.add_run('cronica arretratezza tecnologica')
    run.font.size = Pt(8.8)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run(
        ' della maggior parte delle strutture dilettantistiche locali.'
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'La validità del mercato è confermata dall\'esistenza di concorrenti '
        'già attivi — Playtomic, Wansport, PrenotaWeb — ma la loro copertura '
        'funzionale frammentata e la loro focalizzazione su nicchie specifiche '
        'lasciano aperto uno spazio significativo per un operatore integrato.'
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']

    # Sub-section 1.1
    doc.add_heading('Validazione Preliminare del Mercato', level=2)

    # Stat cards
    stats = [
        ('Gap di Mercato', '65%',
         'ASD italiane prive di sistemi digitali di prenotazione — '
         'opportunità diretta per Sportly'),
        ('Sportivi Attivi', '36%',
         'Praticano sport regolarmente (ISTAT 2024). Utenti abituati '
         'all\'uso di app per i servizi quotidiani'),
        ('Penetrazione Smartphone', '>50%',
         'Adulti italiani connessi. Il canale digitale mobile è già '
         'pronto ad accogliere la domanda'),
    ]

    t = doc.add_table(rows=3, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for col_i, (label, val, desc) in enumerate(stats):
        c = t.cell(0, col_i)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(label.upper())
        run.font.size = Pt(6)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['ink3']
        _cell_shading(c, _hex(C['surface2']))
        _cell_border(c, top=(8, C['c1']), bottom=(4, C['border']),
                     right=(4, C['border']) if col_i < 2 else None)

        c = t.cell(1, col_i)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Space Grotesk'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = [C['c2'], C['c4'], C['c6']][col_i]
        _cell_shading(c, _hex(C['surface2']))
        if col_i < 2:
            _cell_border(c, right=(4, C['border']))

        c = t.cell(2, col_i)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(desc)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink3']
        _cell_shading(c, _hex(C['surface2']))
        if col_i < 2:
            _cell_border(c, right=(4, C['border']))

    _add_spacer(doc, 8)

    # Market data table
    doc.add_heading('Dati di Mercato', level=2)

    t = doc.add_table(rows=7, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    # Header
    for col_i, text in enumerate(['Indicatore di Mercato', 'Dato / Stima']):
        c = t.cell(0, col_i)
        c.text = ''
        _cell_shading(c, _hex(C['ink']))
        p = c.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(6.5)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['white']
        run.font.bold = True

    rows_data = [
        ('Strutture sportive in Italia',
         'Decine di migliaia (palestre, piscine, campi, centri ricreativi)'),
        ('ASD non digitalizzate',
         '~60–70% delle associazioni dilettantistiche attive'),
        ('Popolazione che pratica sport con continuità',
         '~36% degli italiani — Fonte: ISTAT 2024'),
        ('Adulti con smartphone connesso a internet',
         'Oltre il 50% della popolazione adulta italiana'),
        ('Principali concorrenti rilevati',
         'Playtomic, Wansport, PrenotaWeb, Squby'),
        ('Differenziazione rispetto ai concorrenti',
         'Copertura multi-sport + community integrata + dashboard gestori'),
    ]
    for row_i, (label, value) in enumerate(rows_data):
        c = t.cell(row_i + 1, 0)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.color.rgb = C['ink']
        run.font.bold = True

        c = t.cell(row_i + 1, 1)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(8)
        run.font.color.rgb = C['ink2']

        if row_i < 5:
            _cell_border(t.cell(row_i + 1, 0), bottom=(4, C['border']))
            _cell_border(t.cell(row_i + 1, 1), bottom=(4, C['border']))

    _add_spacer(doc, 8)

    # Bar charts
    doc.add_heading('Penetrazione Digitale — Analisi Visiva', level=2)

    bars = [
        ('Strutture non digitalizzate', 'a3c639', '65%'),
        ('Sportivi attivi (ISTAT)', '0d8c8c', '36%'),
        ('Adulti con smartphone', '3a9e7a', '>50%'),
        ('Copertura concorrenti', '2d9a90', '~35%'),
    ]
    for label, color, val in bars:
        _add_colored_bar(doc, label, color, 0.65, val)

    # Callout
    _add_spacer(doc, 4)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    # Bordo sinistro simulato
    run = p.add_run('│ ')
    run.font.size = Pt(12)
    run.font.color.rgb = C['c1']
    run = p.add_run(
        'La coesistenza di concorrenti già operativi e di una quota di mercato '
        'analogica ancora maggioritaria delinea uno scenario ideale per l\'ingresso '
        'di un nuovo operatore: '
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = C['ink2']
    run = p.add_run(
        'la domanda esiste, il mercato è validato, '
        'la soluzione migliore non è ancora stata costruita.'
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = C['ink']

    # Footer
    _add_spacer(doc, 8)
    _add_separator(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        'Sportly — Documento di Progetto — Prima Consegna    ·    '
        '01 Scelta del Progetto    ·    Pagina 02 / 09'
    )
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    # ══════════════════════════════════════
    #  PAGINA 3 — DESCRIZIONE DETTAGLIATA
    # ══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('█ SPORTLY')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run = p.add_run('    02 — Descrizione Dettagliata')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']
    run = p.add_run('                                        03 / 09')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_separator(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SEZIONE 02')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('Descrizione')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run("dell'Idea")
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(26)
    run.font.italic = True
    run.font.color.rgb = C['c2']

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Sportly è un '
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']
    run = p.add_run('ecosistema digitale a due lati')
    run.font.size = Pt(8.8)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run(
        ': da un lato semplifica l\'accesso alle attività sportive per gli '
        'utenti finali, dall\'altro digitalizza la gestione operativa delle '
        'strutture. La piattaforma copre campi sportivi, corsi fitness, '
        'eventi sportivi e attività ricreative — tutto in un\'unica '
        'interfaccia progettata per la massima semplicità d\'uso.'
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']

    # Moduli funzionali
    doc.add_heading('Moduli Funzionali Principali', level=2)

    modules = [
        ('Modulo 01', 'Ricerca Multi-Sport',
         'Filtri per sport, data, distanza e prezzo. Mappa integrata '
         'per trovare strutture nelle vicinanze.', C['c1']),
        ('Modulo 02', 'Calendario Real-Time',
         'Slot disponibili in tempo reale. Prenotazione immediata, '
         'modifica e cancellazione autonome.', C['c2']),
        ('Modulo 03', 'Pagamenti Integrati',
         'Stripe, Apple Pay, Google Pay. Abbonamenti digitali e '
         'ricevute automatiche per strutture e utenti.', C['c6']),
        ('Modulo 04', 'Notifiche Smart',
         'Conferma prenotazione, reminder H-24 e H-2, avvisi slot '
         'liberi, promozioni personalizzate.', C['c7']),
        ('Modulo 05', 'Community e Social',
         'Cerca compagno di gioco, gruppi sportivi locali, badge '
         'fedeltà, classifiche e sfide mensili.', C['c4']),
        ('Modulo 06', 'Dashboard Gestori',
         'Analytics occupazione e ricavi, gestione slot, marketing '
         'diretto, widget booking integrabile.', C['c8']),
    ]

    t = doc.add_table(rows=len(modules), cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for i, (tag, title, desc, color) in enumerate(modules):
        c = t.cell(i, 0)
        c.text = ''
        _cell_border(c, top=(8, color), bottom=(4, C['border']))
        _cell_shading(c, _hex(C['surface2']))

        p = c.paragraphs[0]
        run = p.add_run(tag)
        run.font.size = Pt(6)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = color
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(9)
        run.font.color.rgb = C['ink']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(desc)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink3']

    # Architettura tecnologica
    doc.add_heading('Architettura Tecnologica per Layer', level=2)

    arch = [
        ('FRONTEND', [
            ('React Native',
             'App mobile cross-platform iOS e Android con performance '
             'nativa e codebase unica'),
            ('React.js (Web)',
             'Portale web per gestori e utenti desktop, componenti '
             'condivisi con mobile'),
            ('Google Maps API',
             'Ricerca strutture per geolocalizzazione, navigazione '
             'integrata, mappe interattive'),
        ]),
        ('BACKEND', [
            ('Node.js / Express',
             'API RESTful scalabili, architettura a microservizi '
             'su AWS Lambda'),
            ('Stripe API',
             'Gateway pagamenti PCI-compliant, Apple/Google Pay, '
             'abbonamenti ricorrenti'),
            ('Firebase FCM',
             'Cloud messaging push notifications cross-platform, '
             'alta affidabilità'),
        ]),
        ('DATABASE', [
            ('PostgreSQL',
             'Database relazionale principale per utenti, prenotazioni, '
             'strutture e transazioni'),
            ('Redis Cache',
             'Cache sessioni e slot disponibilità in real-time per '
             'performance ottimale'),
            ('AWS S3',
             'Storage immagini strutture, documenti e asset statici '
             'con CDN globale'),
        ]),
        ('DEVOPS', [
            ('AWS / GCP',
             'Cloud hosting scalabile, autoscaling in base al carico, '
             'multi-region'),
            ('Docker + CI/CD',
             'Containerizzazione e pipeline di rilascio automatizzata '
             'per aggiornamenti rapidi'),
            ('Sentry / Datadog',
             'Monitoraggio errori, performance e uptime in '
             'produzione 24/7'),
        ]),
    ]

    for layer_name, techs in arch:
        t = doc.add_table(rows=1, cols=len(techs))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = True

        for col_i, (name, desc) in enumerate(techs):
            c = t.cell(0, col_i)
            c.text = ''
            _cell_shading(c, _hex(C['surface2']))
            _cell_border(c, top=(8, C['c2']), bottom=(4, C['border']),
                         right=(4, C['border']) if col_i < len(techs) - 1
                         else None)

            p = c.paragraphs[0]
            run = p.add_run(f'[{layer_name}] ')
            run.font.size = Pt(5)
            run.font.name = 'JetBrains Mono'
            run.font.color.rgb = C['ink4']
            p.paragraph_format.space_after = Pt(1)

            p = c.add_paragraph()
            run = p.add_run(name)
            run.font.size = Pt(8)
            run.font.color.rgb = C['ink']
            run.font.bold = True
            p.paragraph_format.space_after = Pt(1)

            p = c.add_paragraph()
            run = p.add_run(desc)
            run.font.size = Pt(7)
            run.font.color.rgb = C['ink3']

        _add_spacer(doc, 3)

    # Footer
    _add_separator(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        'Sportly — Documento di Progetto — Prima Consegna    ·    '
        "02 Descrizione dell'Idea    ·    Pagina 03 / 09"
    )
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    # ══════════════════════════════════════
    #  PAGINA 4 — IL PROBLEMA
    # ══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('█ SPORTLY')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run = p.add_run('    03 — Problema che si Vuole Risolvere')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']
    run = p.add_run('                                   04 / 09')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_separator(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SEZIONE 03')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('Il ')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run('Problema')
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(26)
    run.font.italic = True
    run.font.color.rgb = C['c2']

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Il mercato sportivo dilettantistico italiano opera su infrastrutture '
        'obsolete. Per chi vuole prenotare un campo o un corso, il percorso '
        'tipico è: cercare il numero di telefono della struttura, chiamare '
        'negli orari di apertura, attendere conferma via WhatsApp, pagare in '
        'contanti all\'arrivo. Per i gestori, ogni prenotazione richiede '
        'intervento manuale con elevato rischio di errore. '
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']
    run = p.add_run(
        'Il problema non è culturale — è strutturale.'
    )
    run.font.size = Pt(8.8)
    run.font.bold = True
    run.font.color.rgb = C['ink']

    doc.add_heading('Mappa Pain Point — Soluzione', level=2)

    # Problem-Solution table
    t = doc.add_table(rows=6, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    # Headers
    c = t.cell(0, 0)
    c.text = ''
    _cell_shading(c, _hex(C['red']))
    p = c.paragraphs[0]
    run = p.add_run('PROBLEMI')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['white']
    run.font.bold = True

    c = t.cell(0, 1)
    c.text = ''
    _cell_shading(c, _hex(C['c6']))
    p = c.paragraphs[0]
    run = p.add_run('SOLUZIONI SPORTLY')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['white']
    run.font.bold = True

    problems = [
        ('Prenotazioni ancora manuali',
         'Telefonate, Excel, registri cartacei. Doppie prenotazioni, '
         'errori frequenti e sovraccarico del personale.'),
        ('Nessuna visibilità real-time',
         'Impossibile sapere se un campo è disponibile senza chiamare. '
         'Gli slot liberi si perdono per mancanza di visibilità digitale.'),
        ('Community sportiva frammentata',
         'Nessuno strumento per trovare compagni di gioco, organizzare '
         'tornei amatoriali o aggregare gruppi sportivi locali.'),
        ('Piccole strutture invisibili online',
         'Le ASD locali senza presenza digitale perdono clientela '
         'rispetto a palestre di catena già presenti su Google.'),
        ('Pagamenti e contabilità disorganizzati',
         'Incassi in contanti con rischio ammanchi, abbonamenti tracciati '
         'su carta, zero reportistica automatica.'),
    ]

    solutions = [
        ('Booking digitale 24/7',
         'Calendario smart con disponibilità in tempo reale. '
         'Prenotazioni autonome in 30 secondi da smartphone.'),
        ('Accesso multi-dispositivo',
         'App mobile + portale web. Disponibilità sempre visibile, '
         'prenotabile in qualsiasi momento, da qualsiasi dispositivo.'),
        ('Social feature native',
         'Cerca compagno, gruppi sportivi, notifiche tornei ed eventi '
         'locali. Gamification con badge e classifiche.'),
        ('Marketplace centralizzato',
         'Visibilità uniforme per tutte le strutture indipendentemente '
         'dalle dimensioni. Ogni ASD locale diventa competitiva.'),
        ('Pagamenti e analytics automatizzati',
         'Gateway sicuri PCI-compliant, ricevute digitali, report '
         'occupazione e ricavi in tempo reale.'),
    ]

    for row_i, ((pt, ptx), (st, stx)) in enumerate(
            zip(problems, solutions)):
        # Problema
        c = t.cell(row_i + 1, 0)
        c.text = ''
        _cell_border(c, left=(16, C['red']), bottom=(4, C['border']))
        _cell_shading(c, 'fef5f5')

        p = c.paragraphs[0]
        run = p.add_run(f'Problema 0{row_i + 1}')
        run.font.size = Pt(5.5)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['red']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(pt)
        run.font.size = Pt(8)
        run.font.color.rgb = C['ink']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(ptx)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink3']

        # Soluzione
        c = t.cell(row_i + 1, 1)
        c.text = ''
        _cell_border(c, left=(16, C['c6']), bottom=(4, C['border']))
        _cell_shading(c, 'f0faf5')

        p = c.paragraphs[0]
        run = p.add_run(f'Soluzione 0{row_i + 1}')
        run.font.size = Pt(5.5)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['c6']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(st)
        run.font.size = Pt(8)
        run.font.color.rgb = C['ink']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(stx)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink3']

    # Callout
    _add_spacer(doc, 4)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run('│ ')
    run.font.size = Pt(12)
    run.font.color.rgb = C['c1']
    run = p.add_run(
        'Ogni pain point identificato ha una soluzione tecnologica concreta '
        'e misurabile in Sportly. Il valore non è teorico: è '
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = C['ink2']
    run = p.add_run(
        'operativo, immediato e percepibile'
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run(
        ' sia dagli sportivi che dai gestori delle strutture, dal primo '
        'giorno di utilizzo.'
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = C['ink2']

    _add_separator(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        'Sportly — Documento di Progetto — Prima Consegna    ·    '
        '03 Il Problema    ·    Pagina 04 / 09'
    )
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    # ══════════════════════════════════════
    #  PAGINA 5 — VALORE PER IL CLIENTE
    # ══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('█ SPORTLY')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run = p.add_run('    04 — Valore per il Cliente')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']
    run = p.add_run('                                        05 / 09')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_separator(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SEZIONE 04')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Valore per')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run('il Cliente')
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(26)
    run.font.italic = True
    run.font.color.rgb = C['c2']

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Sportly crea valore concreto e misurabile per entrambi i lati del '
        'marketplace. Per gli utenti sportivi: '
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']
    run = p.add_run('meno attrito, più sport')
    run.font.size = Pt(8.8)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run('. Per i gestori delle strutture: ')
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']
    run = p.add_run(
        'meno telefonate, più prenotazioni, meno errori'
    )
    run.font.size = Pt(8.8)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run(
        '. Il valore non è distribuito uniformemente — è moltiplicato '
        'dall\'effetto rete: più strutture, più utenti. Più utenti, '
        'più strutture.'
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']

    doc.add_heading('Value Proposition per Segmento', level=2)

    # Value grid — Utenti
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('UTENTI SPORTIVI (B2C)')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    utenti_vals = [
        'Risparmio di tempo reale — Prenotazione in meno di 30 secondi da smartphone',
        'Disponibilità continua — Accesso al calendario 24/7',
        'Scelta e confronto trasparente — Prezzo, distanza, recensioni',
        'Flessibilità senza burocrazia — Cancellazioni e modifiche autonome',
        'Community sportiva locale — Trova compagni, entra in gruppi, partecipa a tornei',
        'Motivazione continuativa — Badge fedeltà, punti e sfide mensili',
    ]
    _add_bullet_list(doc, utenti_vals, C['c1'])

    _add_spacer(doc, 4)

    # Value grid — Gestori
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('GESTORI STRUTTURE (B2B)')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    gestori_vals = [
        'Aumento tasso di occupazione — Visibilità su marketplace porta nuovi clienti',
        'Analytics e dati decisionali — Dashboard con occupazione e trend',
        'Automazione operativa — Zero telefonate, zero errori, zero doppie prenotazioni',
        'Pagamenti certi e sicuri — Gateway PCI-compliant, contabilità automatica',
        'Marketing diretto integrato — Promozioni push mirate, senza costi extra',
        'Widget booking integrabile — Sistema prenotazione installabile in pochi minuti',
    ]
    _add_bullet_list(doc, gestori_vals, C['c2'])

    _add_spacer(doc, 4)

    # Copertura competitiva
    doc.add_heading('Copertura Funzionale — Confronto Competitivo', level=2)

    bars = [
        ('Sportly (target)', 'a3c639', '95%'),
        ('Playtomic', '3a9e7a', '60%'),
        ('PrenotaWeb', '2d9a90', '45%'),
        ('Squby', '5daa60', '35%'),
        ('Soluzioni analogiche', 'eaeae4', '— %'),
    ]
    for label, color, val in bars:
        _add_colored_bar(doc, label, color, 0.65, val)

    _add_separator(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        'Sportly — Documento di Progetto — Prima Consegna    ·    '
        '04 Valore per il Cliente    ·    Pagina 05 / 09'
    )
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    # ══════════════════════════════════════
    #  PAGINA 6 — TARGET CLIENTI
    # ══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('█ SPORTLY')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run = p.add_run('    05 — Analisi Target Clienti')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']
    run = p.add_run('                                        06 / 09')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_separator(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SEZIONE 05')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('Target')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run('Clienti')
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(26)
    run.font.italic = True
    run.font.color.rgb = C['c2']

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Sportly si rivolge a tre segmenti principali con esigenze e '
        'comportamenti distinti. La scelta di prioritizzare questi segmenti '
        'nasce dall\'analisi della penetrazione degli smartphone, dalla '
        'propensione all\'uso di app per servizi quotidiani e dalla '
        'frequenza di utilizzo delle strutture sportive.'
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']

    doc.add_heading('Segmentazione Principale', level=2)

    segments = [
        ('SEGMENTO A — PRIMARIO', 'Sportivi Digitali',
         '18–35 anni, praticano sport regolarmente, usano app per quasi tutto.',
         ['Smartphone-native', 'Sport 2–4 volte/settimana',
          'Padel, fitness, running, calcetto',
          'Disposti a pagare per comodità',
          'Influencer nella cerchia sociale'], C['c1']),
        ('SEGMENTO B — SECONDARIO', 'Famiglie Attive',
         '35–50 anni con figli. Cercano corsi per bambini, attività ricreative.',
         ['Utenti smartphone maturi', 'Pianificazione settimanale',
          'Nuoto, tennis, ginnastica bimbi',
          'Sensibili a recensioni e sicurezza',
          'Alta fedeltà una volta acquisiti'], C['c2']),
        ('SEGMENTO C — B2B', 'Gestori ASD',
         'Responsabili di palestre, campi sportivi, centri ricreativi.',
         ['1–50 dipendenti', 'Budget IT limitato',
          'Decisore diretto (titolare/direttore)',
          'Priorità: semplicità di onboarding',
          'Alta retention se onboarding ok'], C['c6']),
    ]

    t = doc.add_table(rows=len(segments), cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for i, (tag, title, desc, traits, color) in enumerate(segments):
        c = t.cell(i, 0)
        c.text = ''
        _cell_border(c, top=(8, color), bottom=(4, C['border']))
        _cell_shading(c, _hex(C['surface2']))

        p = c.paragraphs[0]
        run = p.add_run(tag)
        run.font.size = Pt(6)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = color
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(12)
        run.font.name = 'Instrument Serif'
        run.font.color.rgb = C['ink']
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(desc)
        run.font.size = Pt(7.8)
        run.font.color.rgb = C['ink3']
        p.paragraph_format.space_after = Pt(2)

        for trait in traits:
            p = c.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run('— ')
            run.font.size = Pt(7)
            run.font.color.rgb = C['ink4']
            run = p.add_run(trait)
            run.font.size = Pt(7)
            run.font.name = 'JetBrains Mono'
            run.font.color.rgb = C['ink2']

    _add_spacer(doc, 8)

    doc.add_heading('Comportamento Digitale per Segmento', level=2)

    t = doc.add_table(rows=6, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    headers = ['Dimensione', 'Sportivi Digitali (A)',
               'Famiglie Attive (B)', 'Gestori ASD (C)']
    for col_i, text in enumerate(headers):
        c = t.cell(0, col_i)
        c.text = ''
        _cell_shading(c, _hex(C['ink']))
        p = c.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(6.5)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['white']

    behavior_data = [
        ('Uso smartphone', 'Intensivo, quotidiano',
         'Regolare, funzionale', 'Variabile, spesso limitato'),
        ('Canale acquisizione', 'Social media, passaparola',
         'Google, passaparola', 'Email B2B, fiere sport'),
        ('Frequenza utilizzo app', '2–4 volte/settimana',
         '1–2 volte/settimana', 'Giornaliero (gestione)'),
        ('Sensibilità al prezzo', 'Media — paga per convenienza',
         'Alta — cerca valore', 'Alta — ROI deve essere chiaro'),
        ('Feature più rilevante', 'Booking veloce + community',
         'Affidabilità + semplicità',
         'Dashboard analytics + automazione'),
    ]

    for row_i, (dim, a, b, c_val) in enumerate(behavior_data):
        c = t.cell(row_i + 1, 0)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(dim)
        run.font.size = Pt(8)
        run.font.color.rgb = C['ink']
        run.font.bold = True

        for col_i, val in enumerate([a, b, c_val], 1):
            c = t.cell(row_i + 1, col_i)
            c.text = ''
            p = c.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(7.5)
            run.font.color.rgb = C['ink2']

        if row_i < 4:
            for col_i in range(4):
                _cell_border(t.cell(row_i + 1, col_i),
                             bottom=(4, C['border']))

    _add_spacer(doc, 8)

    doc.add_heading('Dimensionamento del Mercato Accessibile', level=2)

    _add_colored_bar(doc, 'Sportivi 18–35 attivi', 'a3c639', 0.85,
                     '~8M persone')
    _add_colored_bar(doc, 'Famiglie sportive 35–50', '0d8c8c', 0.60,
                     '~5M persone')
    _add_colored_bar(doc, 'ASD target (non digitali)', '3a9e7a', 0.45,
                     '~40K strutture')

    # Callout
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run('│ ')
    run.font.size = Pt(12)
    run.font.color.rgb = C['c1']
    run = p.add_run(
        'La strategia di acquisizione segue un approccio a '
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = C['ink2']
    run = p.add_run('effetto rete progressivo')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run(
        ': si parte dall\'onboarding di strutture partner in una o due '
        'città pilota (es. Milano, Bologna), costruendo una massa critica '
        'di offerta che poi attrae organicamente gli utenti sportivi della zona.'
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = C['ink2']

    _add_separator(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        'Sportly — Documento di Progetto — Prima Consegna    ·    '
        '05 Target Clienti    ·    Pagina 06 / 09'
    )
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    # ══════════════════════════════════════
    #  PAGINA 7 — ANALISI CONCORRENZA
    # ══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('█ SPORTLY')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run = p.add_run('    06 — Analisi della Concorrenza')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']
    run = p.add_run('                                        07 / 09')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_separator(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SEZIONE 06')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Analisi della')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run('Concorrenza')
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(26)
    run.font.italic = True
    run.font.color.rgb = C['c2']

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Il mercato presenta sia concorrenti diretti — piattaforme di '
        'prenotazione sportiva — sia concorrenti indiretti rappresentati '
        'dai metodi analogici ancora dominanti. Nessuno dei player esistenti '
        'copre in modo soddisfacente l\'intero spettro di sport, la '
        'componente community e il supporto completo per i gestori.'
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']

    doc.add_heading('Matrice Competitiva — Feature Comparison', level=2)

    # Competitive table
    t = doc.add_table(rows=7, cols=8)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    comp_headers = [
        'Piattaforma', 'Multi-\nSport', 'Community',
        'Dashboard\nGestori', 'Pagamento\nIntegrato',
        'Notifiche\nSmart', 'App\nMobile', 'Free Tier\nStrutture'
    ]
    for col_i, text in enumerate(comp_headers):
        c = t.cell(0, col_i)
        c.text = ''
        _cell_shading(c, _hex(C['ink']))
        p = c.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(6)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['white']

    comp_data = [
        ('Sportly',      ['S','S','S','S','S','S','S'], True),
        ('Playtomic',    ['P','P','S','S','P','S','N'], False),
        ('Wansport',     ['S','N','P','P','N','P','P'], False),
        ('PrenotaWeb',   ['P','N','S','N','N','N','S'], False),
        ('Squby',        ['P','N','P','N','N','P','N'], False),
        ('Analogico',    ['N','N','N','N','N','N','N'], False),
    ]

    for row_i, (name, checks, is_sportly) in enumerate(comp_data):
        c = t.cell(row_i + 1, 0)
        c.text = ''
        p = c.paragraphs[0]
        if is_sportly:
            _cell_shading(c, 'f5fae0')
            run = p.add_run(name)
            run.font.bold = True
        else:
            run = p.add_run(name)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink']

        for col_i, check in enumerate(checks):
            c = t.cell(row_i + 1, col_i + 1)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(check)
            run.font.size = Pt(7)
            run.font.name = 'JetBrains Mono'
            run.font.bold = True
            if check == 'S':
                run.font.color.rgb = C['c6']
                _cell_shading(c, 'f0faf5')
            elif check == 'P':
                run.font.color.rgb = RGBColor(0x7a, 0x98, 0x20)
                _cell_shading(c, 'f5fae0')
            else:
                run.font.color.rgb = C['ink4']
                _cell_shading(c, _hex(C['surface3']))

            if is_sportly:
                _cell_shading(c, 'f5fae0')

        if row_i < 5:
            for col_i in range(8):
                _cell_border(t.cell(row_i + 1, col_i),
                             bottom=(4, C['border']))

    # Legend
    _add_spacer(doc, 3)
    p = doc.add_paragraph()
    run = p.add_run(
        'S = Sì (completo)   P = Parziale   N = Non presente'
    )
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']

    _add_spacer(doc, 6)

    doc.add_heading('Vantaggi Differenziali', level=2)

    advantages = [
        ('Copertura Multi-Sport Completa',
         'Playtomic copre solo racket sports. Sportly supporta qualsiasi '
         'disciplina — dal nuoto alla yoga, dal calcetto al climbing.'),
        ('Community Integrata Nativa',
         'Nessun concorrente offre funzionalità social native. Sportly '
         'costruisce network effect: più utenti, più valore per tutti.'),
        ('Free Tier per ASD Piccole',
         'Piano gratuito abbassa la barriera all\'adozione per le ASD minori. '
         'Riduce il rischio percepito e accelera la crescita dell\'offerta.'),
        ('Focus sul Mercato Locale Italiano',
         'Concorrenti internazionali non coprono capillarmente il territorio '
         'italiano. Vantaggio di prossimità e conoscenza locale.'),
    ]

    t = doc.add_table(rows=len(advantages), cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for i, (title, desc) in enumerate(advantages):
        c = t.cell(i, 0)
        c.text = ''
        _cell_border(c, top=(8, [C['c1'], C['c2'], C['c6'], C['c7']][i]),
                     bottom=(4, C['border']))
        _cell_shading(c, _hex(C['surface2']))
        p = c.paragraphs[0]
        run = p.add_run(f'Vantaggio 0{i + 1}')
        run.font.size = Pt(6)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = [C['c1'], C['c2'], C['c6'], C['c7']][i]
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)
        p = c.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(8.5)
        run.font.color.rgb = C['ink']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)
        p = c.add_paragraph()
        run = p.add_run(desc)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink3']

    _add_separator(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        'Sportly — Documento di Progetto — Prima Consegna    ·    '
        '06 Analisi Concorrenza    ·    Pagina 07 / 09'
    )
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    # ══════════════════════════════════════
    #  PAGINA 8 — MODELLO DI BUSINESS
    # ══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('█ SPORTLY')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run = p.add_run('    07 — Modello di Business')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']
    run = p.add_run('                                        08 / 09')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_separator(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SEZIONE 07')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Modello di')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run('Business')
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(26)
    run.font.italic = True
    run.font.color.rgb = C['c2']

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Sportly adotta un modello di business '
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']
    run = p.add_run('multi-revenue a tre pilastri')
    run.font.size = Pt(8.8)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    run = p.add_run(
        ': commissioni sulle prenotazioni (scala con il volume), '
        'abbonamenti SaaS per le strutture (ricorrente prevedibile) e '
        'pubblicità/eventi (upsell progressivo). Il modello è progettato '
        'per essere sostenibile fin dal primo anno e scalabile senza '
        'costi marginali significativi.'
    )
    run.font.size = Pt(8.8)
    run.font.color.rgb = C['ink2']

    doc.add_heading('Flussi di Ricavo', level=2)

    revenue_streams = [
        ('01', 'Commissioni Prenotazioni',
         'Percentuale su ogni transazione completata tramite la piattaforma. '
         'Scala direttamente con il volume.',
         '1,5% per transazione completata'),
        ('02', 'SaaS Strutture (B2B)',
         'Abbonamenti mensili ricorrenti per le strutture sportive. '
         'Ricavo prevedibile, bassa churn se onboarding efficace.',
         '29€/mese (Standard) · 59€/mese (Premium)'),
        ('03', 'Pubblicità ed Eventi',
         'Promozione di eventi sportivi, tornei e brand del settore '
         'fitness/sport direttamente sulla piattaforma.',
         'Da 500€/mese (scala con utenti)'),
    ]

    t = doc.add_table(rows=len(revenue_streams), cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for i, (num, title, desc, est) in enumerate(revenue_streams):
        c = t.cell(i, 0)
        c.text = ''
        _cell_shading(c, _hex(C['surface2']))
        _cell_border(c, bottom=(4, C['border']))

        p = c.paragraphs[0]
        run = p.add_run(num)
        run.font.name = 'Space Grotesk'
        run.font.size = Pt(24)
        run.font.color.rgb = C['border']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(f'Revenue Stream {chr(65 + i)}')
        run.font.size = Pt(6)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['c2']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(9)
        run.font.color.rgb = C['ink']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(desc)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink3']
        p.paragraph_format.space_after = Pt(2)

        p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        # Thin top border via table
        run = p.add_run(est)
        run.font.size = Pt(7)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['c1']
        run.font.bold = True

    _add_spacer(doc, 6)

    doc.add_heading('Piani SaaS per Strutture', level=2)

    t = doc.add_table(rows=8, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    pricing = [
        ('Base', '0 €', 'gratuito per sempre',
         ['Fino a 2 risorse/campi', 'Calendario booking base',
          'Notifiche email automatiche', 'Profilo struttura pubblico'],
         False),
        ('Standard', '29 €', 'al mese',
         ['Fino a 10 risorse', 'Analytics base occupazione',
          'Notifiche push personalizzate', 'QR code check-in',
          'Supporto email prioritario'],
         True),
        ('Premium', '59 €', 'al mese',
         ['Risorse illimitate', 'Analytics avanzate + export',
          'Marketing diretto push', 'API + widget integrabile',
          'Supporto telefonico dedicato'],
         False),
    ]

    for col_i, (tier, price, period, features, featured) in enumerate(pricing):
        c = t.cell(0, col_i)
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if featured:
            _cell_shading(c, 'e8f5f5')
            _cell_border(c, top=(8, C['c2']))
            run = p.add_run('CONSIGLIATO')
            run.font.size = Pt(5)
            run.font.name = 'JetBrains Mono'
            run.font.color.rgb = C['white']
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(tier)
            run.font.size = Pt(6.5)
            run.font.name = 'JetBrains Mono'
            run.font.color.rgb = C['ink3']
        else:
            _cell_border(c, top=(4, C['border']))
            run = p.add_run(tier)
            run.font.size = Pt(6.5)
            run.font.name = 'JetBrains Mono'
            run.font.color.rgb = C['ink3']

        c = t.cell(1, col_i)
        c.text = ''
        _cell_shading(c, 'e8f5f5' if featured else _hex(C['surface2']))
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(price)
        run.font.name = 'Space Grotesk'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = C['ink']

        c = t.cell(2, col_i)
        c.text = ''
        _cell_shading(c, 'e8f5f5' if featured else _hex(C['surface2']))
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(period)
        run.font.size = Pt(7)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['ink4']

        for feat_i, feat in enumerate(features):
            c = t.cell(3 + feat_i, col_i)
            c.text = ''
            _cell_shading(c, 'e8f5f5' if featured else _hex(C['surface2']))
            p = c.paragraphs[0]
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run('● ')
            run.font.size = Pt(6)
            run.font.color.rgb = C['c1']
            run = p.add_run(feat)
            run.font.size = Pt(7)
            run.font.color.rgb = C['ink2']

    _add_spacer(doc, 6)

    doc.add_heading('Proiezione Ricavi — Anno 1', level=2)

    t = doc.add_table(rows=6, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    rev_headers = [
        'Fonte di Ricavo', 'Ipotesi Mese 6',
        'Stima €/mese M6', 'Ipotesi Mese 12', 'Stima €/mese M12'
    ]
    for col_i, text in enumerate(rev_headers):
        c = t.cell(0, col_i)
        c.text = ''
        _cell_shading(c, _hex(C['ink']))
        p = c.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(6)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['white']

    rev_data = [
        ('Commissioni (1,5% × 15€)', '1.000 prenotazioni',
         '1.500 €', '5.000 prenotazioni', '7.500 €'),
        ('SaaS Standard (29€/mese)', '30 strutture',
         '870 €', '80 strutture', '2.320 €'),
        ('SaaS Premium (59€/mese)', '10 strutture',
         '590 €', '30 strutture', '1.770 €'),
        ('Pubblicità ed eventi', 'Limitata',
         '200 €', 'Attiva', '1.000 €'),
    ]

    for row_i, (src, hyp6, est6, hyp12, est12) in enumerate(rev_data):
        c = t.cell(row_i + 1, 0)
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(src)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink']
        run.font.bold = True

        for col_i, val in enumerate([hyp6, est6, hyp12, est12], 1):
            c = t.cell(row_i + 1, col_i)
            c.text = ''
            p = c.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(7)
            run.font.name = 'JetBrains Mono'
            if col_i in [2, 4]:
                run.font.color.rgb = C['c6']
                run.font.bold = True
            else:
                run.font.color.rgb = C['ink2']

        if row_i < 3:
            for col_i in range(5):
                _cell_border(t.cell(row_i + 1, col_i),
                             bottom=(4, C['border']))

    # Total row
    c = t.cell(5, 0)
    c.text = ''
    _cell_shading(c, _hex(C['surface2']))
    p = c.paragraphs[0]
    run = p.add_run('Totale Stimato')
    run.font.size = Pt(8)
    run.font.color.rgb = C['ink']
    run.font.bold = True

    for col_i in range(1, 5):
        c = t.cell(5, col_i)
        c.text = ''
        _cell_shading(c, _hex(C['surface2']))
        p = c.paragraphs[0]

    c = t.cell(5, 2)
    p = c.paragraphs[0]
    run = p.add_run('3.160 €')
    run.font.size = Pt(10)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c1']
    run.font.bold = True

    c = t.cell(5, 4)
    p = c.paragraphs[0]
    run = p.add_run('12.590 €')
    run.font.size = Pt(10)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c6']
    run.font.bold = True

    _add_separator(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        'Sportly — Documento di Progetto — Prima Consegna    ·    '
        '07 Modello di Business    ·    Pagina 08 / 09'
    )
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    # ══════════════════════════════════════
    #  PAGINA 9 — WBS E TIMELINE
    # ══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('█ SPORTLY')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink']
    run = p.add_run('    08 — WBS e Timeline')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink3']
    run = p.add_run('                                        09 / 09')
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    _add_separator(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SEZIONE 08')
    run.font.size = Pt(6.5)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['c2']
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('WBS e')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C['ink']
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run('Timeline')
    run.font.name = 'Instrument Serif'
    run.font.size = Pt(26)
    run.font.italic = True
    run.font.color.rgb = C['c2']

    doc.add_heading('Work Breakdown Structure — 8 Fasi', level=2)

    # WBS root node
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    c.text = ''
    _cell_shading(c, _hex(C['ink']))
    p = c.paragraphs[0]
    run = p.add_run('Sportly — Progetto Completo')
    run.font.name = 'Space Grotesk'
    run.font.size = Pt(13)
    run.font.color.rgb = C['white']
    run.font.bold = True
    p = c.add_paragraph()
    run = p.add_run('8 FASI · 20 SETTIMANE · 40+ DELIVERABLE')
    run.font.size = Pt(6)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    _add_spacer(doc, 4)

    phases = [
        ('01', 'Analisi e Pianificazione', [
            '1.1 Ricerca di mercato', '1.2 Requisiti funzionali',
            '1.3 Studio GDPR / legale', '1.4 Analisi concorrenza',
            '1.5 Definizione KPI']),
        ('02', 'Progettazione Architettura', [
            '2.1 Scelta stack tecnologico', '2.2 Diagrammi UML / ER',
            '2.3 Architettura cloud', '2.4 Wireframe schermate',
            '2.5 Prototipo UI/UX']),
        ('03', 'Sviluppo Backend', [
            '3.1 Database PostgreSQL', '3.2 API RESTful',
            '3.3 Integrazione Stripe', '3.4 Notifiche Firebase',
            '3.5 QR code check-in']),
        ('04', 'Sviluppo Frontend', [
            '4.1 App React Native', '4.2 Portale web admin',
            '4.3 UI componenti', '4.4 Integrazione API',
            '4.5 Accessibilità / i18n']),
        ('05', 'Integrazione e Test', [
            '5.1 Integrazione moduli', '5.2 Unit + integration test',
            '5.3 Test sicurezza GDPR', '5.4 Stress test performance',
            '5.5 Test UX utenti reali']),
        ('06', 'Beta Testing', [
            '6.1 Selezione gruppo pilota', '6.2 Rilascio beta chiusa',
            '6.3 Raccolta feedback', '6.4 Bug fixing iterativo',
            '6.5 Validazione metriche']),
        ('07', 'Marketing e Lancio', [
            '7.1 Sito promozionale', '7.2 Campagne social',
            '7.3 Onboarding strutture', '7.4 Materiali comunicazione',
            '7.5 Lancio App/Play Store']),
        ('08', 'Monitoraggio Continuo', [
            '8.1 KPI mensili analytics', '8.2 Customer support',
            '8.3 Pianificazione release', '8.4 Manutenzione sicurezza']),
    ]

    phase_colors = [C['c1'], C['c2'], C['c6'], C['c7'],
                    C['c4'], C['c5'], C['c8'], C['c3']]

    for row_i in range(0, len(phases), 4):
        batch = phases[row_i:row_i + 4]
        t = doc.add_table(rows=2, cols=len(batch))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = True

        for col_i, (num, title, items) in enumerate(batch):
            phase_i = row_i + col_i
            color = phase_colors[phase_i]

            # Header
            c = t.cell(0, col_i)
            c.text = ''
            _cell_shading(c, _hex(C['surface3']))
            _cell_border(c, top=(8, color), bottom=(4, C['border']))
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(num)
            run.font.size = Pt(7.5)
            run.font.name = 'JetBrains Mono'
            run.font.color.rgb = color
            run.font.bold = True
            p = c.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.size = Pt(8)
            run.font.color.rgb = C['ink']
            run.font.bold = True

            # Items
            c = t.cell(1, col_i)
            c.text = ''
            _cell_border(c, bottom=(4, C['border']))
            for item in items:
                p = c.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.3)
                code, text = item.split(' ', 1)
                run = p.add_run(code + ' ')
                run.font.size = Pt(6)
                run.font.name = 'JetBrains Mono'
                run.font.color.rgb = color
                run = p.add_run(text)
                run.font.size = Pt(7.5)
                run.font.color.rgb = C['ink2']

        _add_spacer(doc, 4)

    # Timeline / Milestones
    doc.add_heading('Milestone Principali', level=2)

    milestones = [
        ('M1', 'Analisi e Definizione Completata',
         'Fine settimana 2',
         'Documento di progetto — Prima consegna (questo documento)',
         C['c1']),
        ('M2', 'Architettura e Prototipo UI Approvati',
         'Fine settimana 5',
         'Wireframe navigabili + schema database + documentazione stack',
         C['c2']),
        ('M3', 'Backend API Funzionanti e Documentate',
         'Fine settimana 9',
         'API RESTful testate, integrazione Stripe e Firebase FCM attive',
         C['c6']),
        ('M4', 'App Mobile v0.9 — Build Interna Testabile',
         'Fine settimana 13',
         'APK e .ipa distribuibili internamente, flussi principali funzionanti',
         C['c7']),
        ('M5', 'Beta Test Completato e Report Feedback',
         'Fine settimana 17',
         'Report qualitativo e quantitativo dal gruppo pilota',
         C['c4']),
        ('LANCIO', 'Lancio Ufficiale su App Store e Google Play',
         'Fine settimana 20',
         'App pubblicata, prime strutture onboarded, evento lancio',
         C['ink']),
    ]

    t = doc.add_table(rows=len(milestones), cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    for i, (dot, title, week, deliverable, color) in enumerate(milestones):
        c = t.cell(i, 0)
        c.text = ''
        _cell_border(c, left=(8, color), bottom=(4, C['border']))
        _cell_shading(c, _hex(C['surface2']))

        p = c.paragraphs[0]
        run = p.add_run(f'[{dot}]')
        run.font.size = Pt(8)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = color
        run.font.bold = True
        run = p.add_run(f'  {title}')
        run.font.size = Pt(9)
        run.font.color.rgb = C['ink']
        run.font.bold = True
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(week)
        run.font.size = Pt(6.5)
        run.font.name = 'JetBrains Mono'
        run.font.color.rgb = C['c2']
        p.paragraph_format.space_after = Pt(1)

        p = c.add_paragraph()
        run = p.add_run(deliverable)
        run.font.size = Pt(7.5)
        run.font.color.rgb = C['ink3']

    # Closing
    _add_spacer(doc, 12)
    _add_separator(doc)
    _add_spacer(doc, 4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('— Fine del Documento —')
    run.font.size = Pt(9)
    run.font.name = 'Instrument Serif'
    run.font.italic = True
    run.font.color.rgb = C['ink3']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Sportly — Documento di Progetto — Prima Consegna\n'
        'Team 6-3-14 · Gestione Progetto e Organizzazione di Impresa\n'
        'Classe 5I · 10 Aprile 2025'
    )
    run.font.size = Pt(7)
    run.font.name = 'JetBrains Mono'
    run.font.color.rgb = C['ink4']

    # ── Salva ──
    doc.save('Sportly_Prima_Consegna.docx')
    print('✅ Documento salvato: Sportly_Prima_Consegna.docx')


# ═══════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════
if __name__ == '__main__':
    create_sportly_docx()